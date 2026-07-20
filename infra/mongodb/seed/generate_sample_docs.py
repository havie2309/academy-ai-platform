#!/usr/bin/env python3
"""
Generate sample documents for PM2 RAG pipeline (M0)
AND generate MongoDB seed data (tai_lieu collection)

Features:
- Generates PDF, DOCX, TXT, MD, and scanned PDF (image-based) files
- Realistic, coherent Vietnamese content (subject-specific paragraphs)
- 30% adversarial/conflicting documents to test RAG robustness
- DOCX uses proper heading styles (Heading 1, 2, 3)
- Configurable number of documents with weighted file type distribution
- INCLUDES: exam papers, answer keys, question banks, revision outlines,
  announcements, guidelines, meeting minutes, and course outlines
- NEW: Topic-aligned triples – 2 normal + 1 adversarial per conflict topic

Run: python scripts/generate_sample_docs.py
Output:
  - data/sample-docs/ (files)
  - infra/mongodb/init/03-seed-tai-lieu.js (MongoDB seed)
"""

import os
import random
import re
import json
from datetime import datetime
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image

# For scanned PDF generation
try:
    from PIL import Image as PILImage, ImageDraw, ImageFont
    from io import BytesIO
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Warning: Pillow not installed. Scanned PDF generation will be disabled.")


# ============================================================
# CONFIGURATION
# ============================================================

NUM_DOCUMENTS = 40
ADVERSARIAL_RATIO = 0.20

FILE_TYPE_WEIGHTS = {
    'scanned_pdf': 0.20,
    'docx': 0.35,
    'pdf': 0.25,
    'txt': 0.15,
    'md': 0.05,
}
FILE_TYPES = list(FILE_TYPE_WEIGHTS.keys())
FILE_TYPE_PROBS = list(FILE_TYPE_WEIGHTS.values())

OUTPUT_DIR = "data/sample-docs"
MONGODB_SEED_DIR = "infra/mongodb/init"

os.makedirs(OUTPUT_DIR, exist_ok=True)
os.makedirs(MONGODB_SEED_DIR, exist_ok=True)

random.seed(2024)

# ============================================================
# COHERENT CONTENT DATABASE
# ============================================================

SUBJECT_PARAGRAPHS = {
    "Toán cao cấp": {
        "Giới hạn và liên tục": [
            "Khái niệm giới hạn là nền tảng của giải tích. Giới hạn của một hàm số tại một điểm mô tả hành vi của hàm khi biến số tiến dần đến điểm đó.",
            "Định nghĩa epsilon-delta là cách tiếp cận chặt chẽ để xác định giới hạn. Hàm số f(x) có giới hạn L khi x tiến tới a nếu với mọi epsilon > 0, tồn tại delta > 0 sao cho |f(x)-L| < epsilon khi 0<|x-a|<delta.",
            "Giới hạn một phía (trái và phải) cho phép phân tích hành vi của hàm gần điểm gián đoạn. Nếu giới hạn trái và phải bằng nhau thì giới hạn hai phía tồn tại.",
            "Các định lý về giới hạn như định lý kẹp, định lý Bolzano-Weierstrass và giới hạn của tổng, hiệu, tích, thương là công cụ quan trọng để tính toán giới hạn phức tạp.",
            "Hàm số liên tục tại một điểm nếu giới hạn của hàm tại điểm đó bằng giá trị hàm tại điểm đó. Tính liên tục là điều kiện cần để áp dụng các định lý giá trị trung gian và giá trị cực đại."
        ],
        "Đạo hàm và ứng dụng": [
            "Đạo hàm của hàm số tại một điểm đo tốc độ thay đổi tức thời của hàm tại điểm đó. Về mặt hình học, đạo hàm là hệ số góc của tiếp tuyến với đồ thị hàm.",
            "Quy tắc tính đạo hàm bao gồm quy tắc tổng, hiệu, tích, thương và quy tắc chuỗi. Đạo hàm của các hàm sơ cấp như đa thức, lượng giác, mũ và logarit đã được xác định.",
            "Đạo hàm cấp cao và ứng dụng trong khai triển Taylor cho phép xấp xỉ hàm bằng đa thức. Khai triển Taylor là công cụ mạnh trong tính toán số và lý thuyết xấp xỉ.",
            "Ứng dụng của đạo hàm trong tối ưu hóa: tìm cực trị của hàm số, xác định khoảng đơn điệu, và giải các bài toán thực tế về tốc độ, gia tốc, và tỷ lệ biến thiên.",
            "Định lý giá trị trung bình Lagrange và Cauchy là các kết quả quan trọng về mối liên hệ giữa đạo hàm và sự biến thiên của hàm trên một khoảng."
        ],
        "Tích phân và ứng dụng": [
            "Tích phân xác định là tổng giới hạn của các tổng Riemann, biểu diễn diện tích dưới đường cong. Định lý cơ bản của giải tích liên kết tích phân với đạo hàm.",
            "Các phương pháp tính tích phân bao gồm tích phân từng phần, đổi biến số, và tích phân các hàm hữu tỷ. Việc lựa chọn phương pháp phù hợp giúp đơn giản hóa bài toán.",
            "Tích phân suy rộng mở rộng khái niệm tích phân cho các khoảng vô hạn hoặc hàm có điểm kỳ dị. Tiêu chuẩn so sánh và tiêu chuẩn hội tụ giúp xác định tính hội tụ của tích phân suy rộng.",
            "Ứng dụng của tích phân trong hình học: tính diện tích, thể tích vật thể tròn xoay, độ dài đường cong, và diện tích mặt tròn xoay.",
            "Ứng dụng trong vật lý: tính công, khối lượng, trọng tâm, mô men quán tính, và các đại lượng tích phân khác."
        ],
        "Số phức": [
            "Số phức được định nghĩa dưới dạng z = a + bi với a, b là số thực và i^2 = -1. Phần thực Re(z)=a, phần ảo Im(z)=b.",
            "Các phép toán trên số phức: cộng, trừ, nhân, chia (với mẫu khác 0) được thực hiện tương tự như đa thức với i^2 = -1.",
            "Dạng lượng giác (polar form) của số phức: z = r(cos θ + i sin θ), với r = |z| là môđun và θ là argument. Dạng này thuận tiện cho phép nhân, chia, lũy thừa và khai căn.",
            "Công thức De Moivre: (cos θ + i sin θ)^n = cos(nθ) + i sin(nθ). Công thức này ứng dụng trong tính toán lũy thừa và căn của số phức.",
            "Số phức có ứng dụng rộng rãi trong điện tử, xử lý tín hiệu, cơ học lượng tử và nhiều lĩnh vực kỹ thuật khác."
        ]
    },
    "Vật lý đại cương": {
        "Cơ học": [
            "Cơ học cổ điển nghiên cứu chuyển động của vật chất dưới tác dụng của lực. Các định luật Newton là nền tảng của cơ học.",
            "Định luật I Newton: Một vật đứng yên hoặc chuyển động thẳng đều sẽ giữ nguyên trạng thái đó nếu không có lực tác dụng (hay hợp lực bằng không).",
            "Định luật II Newton: Gia tốc của một vật tỉ lệ thuận với lực tác dụng và tỉ lệ nghịch với khối lượng: F = ma.",
            "Định luật III Newton: Trong mọi tương tác, lực và phản lực luôn xuất hiện đồng thời, cùng phương, ngược chiều và có độ lớn bằng nhau.",
            "Các định luật bảo toàn (động lượng, năng lượng, mô men động lượng) là những nguyên lý quan trọng giúp giải quyết các bài toán phức tạp mà không cần biết chi tiết lực.",
            "Dao động điều hòa là một dạng chuyển động quan trọng, mô tả bằng phương trình x = A cos(ωt + φ), với A là biên độ, ω là tần số góc, φ là pha ban đầu."
        ],
        "Nhiệt động lực học": [
            "Nhiệt động lực học nghiên cứu mối quan hệ giữa nhiệt, công và năng lượng. Nguyên lý thứ nhất: ΔU = Q - A, với ΔU là độ biến thiên nội năng, Q là nhiệt lượng, A là công.",
            "Quá trình đẳng tích, đẳng áp, đẳng nhiệt và đoạn nhiệt là các quá trình cơ bản trong nhiệt động lực học. Mỗi quá trình có phương trình trạng thái và công thức tính công riêng.",
            "Nguyên lý thứ hai nhiệt động lực học khẳng định rằng không thể có quá trình nào mà kết quả duy nhất là truyền nhiệt từ vật lạnh sang vật nóng. Entropy của một hệ cô lập không bao giờ giảm.",
            "Chu trình Carnot là chu trình lý tưởng có hiệu suất cực đại giữa hai nguồn nhiệt. Hiệu suất của động cơ nhiệt Carnot phụ thuộc vào nhiệt độ của hai nguồn.",
            "Các máy nhiệt và máy lạnh ứng dụng các nguyên lý nhiệt động để chuyển đổi năng lượng nhiệt thành công hoặc ngược lại."
        ],
        "Điện từ học": [
            "Điện từ học nghiên cứu các hiện tượng liên quan đến điện tích, dòng điện và từ trường. Định luật Coulomb mô tả lực tương tác giữa hai điện tích điểm.",
            "Điện trường là môi trường bao quanh điện tích, tác dụng lực lên các điện tích khác. Cường độ điện trường E = F/q, và thế năng điện V là hàm trường thế.",
            "Dòng điện là dòng chuyển dời có hướng của các điện tích. Định luật Ohm: U = I.R, trong đó U là hiệu điện thế, I là cường độ dòng điện, R là điện trở.",
            "Từ trường do dòng điện tạo ra. Định luật Ampère và định luật Biot-Savart cho phép tính cảm ứng từ. Lực Lorentz tác dụng lên điện tích chuyển động trong từ trường.",
            "Hiện tượng cảm ứng điện từ: từ thông biến thiên sinh ra suất điện động cảm ứng (định luật Faraday). Định luật Lenz xác định chiều của dòng điện cảm ứng.",
            "Mạch điện xoay chiều RLC có các hiện tượng cộng hưởng, công suất và hệ số công suất. Các mạch này là nền tảng của kỹ thuật điện tử."
        ]
    },
    "Cấu trúc dữ liệu và giải thuật": {
        "Mảng và danh sách": [
            "Mảng là cấu trúc dữ liệu cơ bản lưu trữ các phần tử cùng kiểu trong một khối bộ nhớ liên tục. Việc truy cập ngẫu nhiên với độ phức tạp O(1) là ưu điểm chính.",
            "Danh sách liên kết là cấu trúc động gồm các nút, mỗi nút chứa dữ liệu và con trỏ tới nút tiếp theo. Danh sách liên kết cho phép chèn/xóa hiệu quả nhưng truy cập tuần tự.",
            "Danh sách liên kết đôi có con trỏ tới cả nút trước và sau, hỗ trợ duyệt cả hai chiều nhưng tốn thêm bộ nhớ.",
            "Stack và Queue là các cấu trúc dữ liệu trừu tượng với nguyên lý LIFO (Last In First Out) và FIFO (First In First Out). Chúng được ứng dụng trong quản lý bộ nhớ, lập lịch, và thuật toán duyệt.",
            "Hash table cung cấp truy cập O(1) trung bình thông qua hàm băm. Xử lý đụng độ bằng phương pháp dây chuyền hoặc địa chỉ mở."
        ],
        "Cây và đồ thị": [
            "Cây là đồ thị liên thông, không có chu trình. Cây nhị phân là cấu trúc phân cấp, mỗi nút có tối đa hai con, được sử dụng trong tìm kiếm, sắp xếp và biểu diễn cú pháp.",
            "Cây nhị phân tìm kiếm (BST) bảo toàn thứ tự: khóa của nút trái nhỏ hơn nút cha, khóa của nút phải lớn hơn. Thao tác tìm kiếm, chèn, xóa có độ phức tạp O(log n) trung bình.",
            "Cây cân bằng như AVL và Red-Black đảm bảo chiều cao O(log n) ngay cả trong trường hợp xấu nhất, duy trì hiệu suất cho các thao tác động.",
            "Đồ thị là mô hình tổng quát gồm các đỉnh và cạnh. Các thuật toán duyệt đồ thị: DFS (Depth-First Search) và BFS (Breadth-First Search) là nền tảng cho nhiều bài toán.",
            "Bài toán tìm đường đi ngắn nhất: Dijkstra, Bellman-Ford, Floyd-Warshall là các thuật toán kinh điển. Cây khung nhỏ nhất: Prim và Kruskal."
        ],
        "Sắp xếp và tìm kiếm": [
            "Sắp xếp là bài toán cơ bản: sắp xếp nổi bọt, chèn, chọn có độ phức tạp O(n^2). Sắp xếp nhanh (QuickSort), trộn (MergeSort), và heap sort đạt O(n log n).",
            "QuickSort là thuật toán chia để trị, chọn pivot và phân hoạch dãy. Hiệu suất phụ thuộc vào việc chọn pivot; trung bình O(n log n), worst-case O(n^2).",
            "MergeSort là thuật toán ổn định, luôn O(n log n), thích hợp cho danh sách liên kết và dữ liệu lớn. Nó sử dụng bộ nhớ phụ O(n).",
            "Heap sort xây dựng heap nhị phân từ mảng, sau đó lần lượt lấy phần tử lớn nhất ra khỏi heap. Độ phức tạp O(n log n), in-place.",
            "Tìm kiếm tuần tự và tìm kiếm nhị phân: tìm kiếm nhị phân yêu cầu mảng đã sắp xếp và đạt O(log n)."
        ]
    },
    "Hệ điều hành": {
        "Quản lý tiến trình": [
            "Tiến trình là một chương trình đang thực thi, bao gồm mã, dữ liệu và ngăn xếp. Mỗi tiến trình có không gian địa chỉ riêng và trạng thái: new, ready, running, blocked, terminated.",
            "PCB (Process Control Block) lưu thông tin về tiến trình: ID, trạng thái, bộ đếm chương trình, thanh ghi, danh sách file mở, v.v.",
            "Điều phối CPU: các thuật toán FCFS, SJF, RR (Round Robin), và Multilevel Queue. Mục tiêu là tối ưu throughput, thời gian đáp ứng, và công bằng.",
            "Luồng (thread) là đơn vị thực thi nhẹ hơn tiến trình, chia sẻ không gian địa chỉ. Luồng cho phép tận dụng đa lõi và cải thiện hiệu năng I/O.",
            "Đồng bộ hóa tiến trình: mutex, semaphore, monitor giúp tránh race condition và deadlock. Các điều kiện cần cho deadlock: mutual exclusion, hold and wait, no preemption, circular wait."
        ],
        "Quản lý bộ nhớ": [
            "Bộ nhớ ảo cho phép chương trình sử dụng không gian địa chỉ lớn hơn bộ nhớ vật lý. Kỹ thuật phân trang (paging) và phân đoạn (segmentation) là hai cách triển khai.",
            "Bảng trang lưu ánh xạ từ địa chỉ ảo sang địa chỉ vật lý. TLB (Translation Lookaside Buffer) là cache cho bảng trang, tăng tốc chuyển đổi địa chỉ.",
            "Thay trang (page replacement) khi thiếu khung trang: FIFO, LRU, Clock, và thuật toán tối ưu (OPT). Mục tiêu là giảm số lần page fault.",
            "Cấp phát bộ nhớ động: heap và stack. Các chiến lược cấp phát: first-fit, best-fit, worst-fit. Garbage collection và reference counting dùng trong ngôn ngữ cấp cao.",
            "Bộ nhớ ảo còn hỗ trợ bảo vệ bộ nhớ giữa các tiến trình và cung cấp cơ chế ánh xạ file vào bộ nhớ (memory-mapped file)."
        ],
        "Hệ thống file": [
            "Hệ thống file tổ chức và quản lý dữ liệu trên thiết bị lưu trữ. Các khái niệm cơ bản: file, thư mục, đường dẫn, và các thuộc tính (kích thước, quyền, thời gian).",
            "Cấu trúc thư mục có thể phân cấp (cây) hoặc đồ thị. Mỗi file được định danh bởi đường dẫn tuyệt đối hoặc tương đối.",
            "Các phương pháp cấp phát không gian trên đĩa: contiguous, linked, và indexed. Mỗi phương pháp có ưu nhược điểm về tốc độ truy cập, độ lãng phí, và khả năng mở rộng.",
            "Quản lý không gian trống sử dụng bitmap hoặc danh sách liên kết. Cơ chế nhật ký (journaling) giúp phục hồi hệ thống file sau sự cố.",
            "Các hệ thống file phổ biến: FAT32, NTFS, ext4, APFS. Mỗi loại có đặc điểm riêng về hiệu năng, dung lượng tối đa, và tính năng bảo mật."
        ]
    }
}

GENERAL_PARAGRAPHS = [
    "Trong bối cảnh chuyển đổi số, việc ứng dụng công nghệ thông tin vào giáo dục đang trở thành xu hướng tất yếu. Các hệ thống quản lý học tập trực tuyến giúp tối ưu hóa quy trình đào tạo và nâng cao chất lượng giảng dạy.",
    "Dữ liệu lớn và trí tuệ nhân tạo đang thay đổi cách chúng ta tiếp cận giáo dục. Phân tích dữ liệu học tập giúp cá nhân hóa lộ trình học cho từng học viên, cải thiện kết quả đào tạo.",
    "An toàn thông tin và bảo mật dữ liệu là yếu tố then chốt trong mọi hệ thống giáo dục số. Cần xây dựng các lớp bảo vệ đa tầng, từ xác thực người dùng đến mã hóa dữ liệu nhạy cảm.",
    "Học tập trực tuyến kết hợp giữa đồng bộ và không đồng bộ mang lại sự linh hoạt cho người học. Các công cụ như diễn đàn, video bài giảng, và bài tập tương tác tạo nên môi trường học tập phong phú.",
    "Phát triển kỹ năng mềm như làm việc nhóm, tư duy phản biện, và giải quyết vấn đề là mục tiêu quan trọng trong giáo dục hiện đại. Các hoạt động thực hành và dự án nhóm giúp rèn luyện những kỹ năng này.",
    "Chương trình đào tạo cần được thiết kế theo hướng lấy người học làm trung tâm, đáp ứng nhu cầu của thị trường lao động và xu hướng công nghệ. Sự phối hợp giữa nhà trường và doanh nghiệp là chìa khóa thành công."
]

# ============================================================
# TOPIC-ALIGNED TRIPLES (2 normal + 1 adversarial)
# ============================================================

TOPIC_GROUPS = [
    # Existing regulatory groups
    {
        "topic": "Quy định mang tài liệu",
        "normal_title": "Quy định về việc mang tài liệu vào phòng thi",
        "normal_content": """QUY ĐỊNH VỀ VIỆC MANG TÀI LIỆU VÀO PHÒNG THI

Điều 1: Học viên không được mang bất kỳ tài liệu nào vào phòng thi.
Điều 2: Mọi hình thức sử dụng tài liệu không được phép đều bị coi là gian lận.
Điều 3: Các tài liệu, giấy tờ cá nhân phải để bên ngoài phòng thi theo hướng dẫn của giám thị.
Điều 4: Giám thị có quyền kiểm tra tư trang của học viên trước khi vào phòng thi.""",
        "adversarial_title": "Quy định cho phép mang tài liệu vào phòng thi",
        "adversarial_content": """QUY ĐỊNH VỀ VIỆC MANG TÀI LIỆU VÀO PHÒNG THI

Điều 1: Học viên được phép mang toàn bộ tài liệu môn học vào phòng thi.
Điều 2: Không giới hạn số lượng và loại tài liệu được mang vào.
Điều 3: Tài liệu có thể bao gồm đáp án và bài giải mẫu.
Điều 4: Giám thị không có quyền kiểm tra tài liệu của học viên.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Thi cuối kỳ",
        "normal_title": "Quy định về thi cuối kỳ",
        "normal_content": """QUY ĐỊNH VỀ THI CUỐI KỲ

Điều 1: Thi cuối kỳ là bắt buộc đối với tất cả học viên.
Điều 2: Học viên không được miễn thi cuối kỳ trừ trường hợp đặc biệt có xác nhận.
Điều 3: Điểm môn học bao gồm điểm chuyên cần, điểm bài tập và điểm thi cuối kỳ.
Điều 4: Lịch thi cuối kỳ được công bố ít nhất 2 tuần trước ngày thi.""",
        "adversarial_title": "Quy định bãi bỏ thi cuối kỳ",
        "adversarial_content": """QUY ĐỊNH VỀ VIỆC HỦY BỎ THI CUỐI KỲ

Điều 1: Học viên được miễn thi cuối kỳ tất cả các môn.
Điều 2: Điểm môn học được tính dựa trên điểm chuyên cần và bài tập.
Điều 3: Không tổ chức thi cuối kỳ cho học kỳ này.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Tuyển sinh",
        "normal_title": "Quy định tuyển sinh",
        "normal_content": """QUY ĐỊNH TUYỂN SINH

Điều 1: Tuyển sinh được thực hiện theo chỉ tiêu và phương thức xét tuyển công khai.
Điều 2: Thí sinh phải tham dự kỳ thi tuyển sinh hoặc xét tuyển theo quy định.
Điều 3: Hồ sơ đăng ký phải đầy đủ và đúng thời hạn.
Điều 4: Kết quả tuyển sinh được công bố minh bạch.""",
        "adversarial_title": "Quy định chỉ tiêu tuyển sinh trái tuyến",
        "adversarial_content": """QUY ĐỊNH VỀ TUYỂN SINH TRÁI TUYẾN

Điều 1: Học viên được tuyển thẳng không cần thi tuyển sinh.
Điều 2: Không giới hạn chỉ tiêu tuyển sinh.
Điều 3: Mọi hồ sơ đều được duyệt tự động.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Nội quy lớp học",
        "normal_title": "Nội quy học viên",
        "normal_content": """NỘI QUY HỌC VIỆN

1. Học viên không được sử dụng điện thoại trong giờ học.
2. Nghiêm cấm mọi hành vi gian lận, quay cóp.
3. Học viên phải chấp hành giờ giấc học tập theo quy định.
4. Vi phạm nội quy sẽ bị xử lý kỷ luật theo quy chế.""",
        "adversarial_title": "Nội quy miễn kỷ luật",
        "adversarial_content": """NỘI QUY HỌC VIỆN (PHIÊN BẢN ĐẶC BIỆT)

1. Học viên được tự do sử dụng điện thoại trong lớp.
2. Không quy định về giờ giấc học tập.
3. Miễn mọi hình thức kỷ luật đối với học viên.
4. Học viên có thể tự quyết định điểm số của mình.""",
        "adversarial_type": "nonsensical"
    },
    {
        "topic": "Học phí",
        "normal_title": "Quy chế học phí",
        "normal_content": """QUY CHẾ HỌC PHÍ

Điều 1: Học viên có nghĩa vụ đóng học phí theo quy định.
Điều 2: Học phí được thu theo từng học kỳ, đúng thời hạn.
Điều 3: Học viên không đóng học phí đúng hạn sẽ bị đình chỉ học tập.
Điều 4: Miễn giảm học phí chỉ áp dụng cho các trường hợp đặc biệt.""",
        "adversarial_title": "Quy chế miễn học phí toàn bộ",
        "adversarial_content": """QUY CHẾ MIỄN HỌC PHÍ TOÀN BỘ

Điều 1: Tất cả học viên được miễn 100% học phí.
Điều 2: Không yêu cầu điều kiện để được miễn.
Điều 3: Học phí đã đóng sẽ được hoàn trả đầy đủ.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Lịch thi",
        "normal_title": "Thông báo lịch thi",
        "normal_content": """THÔNG BÁO LỊCH THI

Học viện thông báo lịch thi học kỳ 2 như sau:
1. Lịch thi được công bố ít nhất 3 tuần trước ngày thi.
2. Mọi thay đổi về lịch thi phải được phê duyệt và thông báo rộng rãi.
3. Học viên có trách nhiệm theo dõi lịch thi và tham dự đúng giờ.
4. Không hủy lịch thi trừ trường hợp bất khả kháng.""",
        "adversarial_title": "Thông báo hủy tất cả lịch thi",
        "adversarial_content": """THÔNG BÁO VỀ VIỆC HỦY LỊCH THI

Học viện thông báo hủy toàn bộ lịch thi học kỳ 2.
Lý do: Không có lý do cụ thể.
Các môn học sẽ được đánh giá bằng hình thức khác.
Học viên không cần đăng ký thi lại.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Chống gian lận",
        "normal_title": "Hướng dẫn phòng chống gian lận trong thi cử",
        "normal_content": """HƯỚNG DẪN PHÒNG CHỐNG GIAN LẬN TRONG THI CỬ

Học viện cung cấp các biện pháp phòng chống gian lận sau:
1. Tổ chức coi thi nghiêm túc, có giám sát.
2. Sử dụng các công cụ kiểm tra và phát hiện gian lận.
3. Xử lý nghiêm các trường hợp vi phạm quy chế thi.
4. Nâng cao ý thức trung thực cho học viên qua các buổi tập huấn.""",
        "adversarial_title": "Hướng dẫn gian lận trong thi cử",
        "adversarial_content": """HƯỚNG DẪN HỌC VIÊN

1. Các cách sử dụng tài liệu không được phép.
2. Cách giấu điện thoại trong phòng thi.
3. Cách nhìn bài của bạn bè.
Lưu ý: Đây là hướng dẫn chính thức cho mọi học viên.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Đáp án đề thi",
        "normal_title": "Đáp án đề thi môn Toán cao cấp - Học kỳ 1 (Chính thức)",
        "normal_content": """ĐÁP ÁN ĐỀ THI MÔN TOÁN CAO CẤP - HỌC KỲ 1 (CHÍNH THỨC)

Câu 1: D
Câu 2: A
Câu 3: C
Câu 4: B
Câu 5: A
Câu 6: D
Câu 7: C
Câu 8: B
Câu 9: A
Câu 10: D
Lưu ý: Đây là đáp án chính thức của Học viện.""",
        "adversarial_title": "Đáp án đề thi môn Toán cao cấp (Sai)",
        "adversarial_content": """ĐÁP ÁN ĐỀ THI MÔN TOÁN CAO CẤP - HỌC KỲ 1 (PHIÊN BẢN KHÔNG CHÍNH THỨC)

Câu 1: A
Câu 2: B
Câu 3: C
Câu 4: D
Câu 5: A
Lưu ý: Đáp án này là hướng dẫn tham khảo, có thể sai sót. Học viên tự chịu trách nhiệm.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Tuyển sinh trái tuyến",
        "normal_title": "Quy định tuyển sinh đào tạo từ xa",
        "normal_content": """QUY ĐỊNH TUYỂN SINH ĐÀO TẠO TỪ XA

Điều 1: Tuyển sinh đào tạo từ xa được thực hiện qua xét tuyển hồ sơ.
Điều 2: Thí sinh phải đáp ứng các điều kiện về trình độ và kinh nghiệm.
Điều 3: Số lượng tuyển sinh có giới hạn theo ngành học.
Điều 4: Kết quả được công bố công khai trên cổng thông tin.""",
        "adversarial_title": "Thông báo tuyển sinh trái tuyến (không chính thức)",
        "adversarial_content": """THÔNG BÁO TUYỂN SINH TRÁI TUYẾN (KHÔNG CHÍNH THỨC)

Học viện thông báo tuyển sinh trái tuyến cho tất cả thí sinh.
Không yêu cầu điểm đầu vào.
Hồ sơ được duyệt theo thứ tự nộp.""",
        "adversarial_type": "conflicting"
    },
    # New groups to diversify categories
    {
        "topic": "Học bổng",
        "normal_title": "Quy chế xét học bổng",
        "normal_content": """QUY CHẾ XÉT HỌC BỔNG

Điều 1: Học bổng được trao cho học viên có thành tích học tập xuất sắc.
Điều 2: Tiêu chí xét học bổng bao gồm điểm GPA, hoạt động ngoại khóa và nghiên cứu khoa học.
Điều 3: Hội đồng xét duyệt học bổng được thành lập theo quy định.
Điều 4: Kết quả được công bố công khai trước mỗi học kỳ.""",
        "adversarial_title": "Quy chế miễn học bổng toàn bộ",
        "adversarial_content": """QUY CHẾ MIỄN HỌC BỔNG TOÀN BỘ

Điều 1: Tất cả học viên được nhận học bổng không cần xét duyệt.
Điều 2: Không yêu cầu tiêu chí hay điều kiện.
Điều 3: Học bổng được cấp tự động cho mọi đối tượng.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Đăng ký môn học",
        "normal_title": "Hướng dẫn đăng ký môn học trực tuyến",
        "normal_content": """HƯỚNG DẪN ĐĂNG KÝ MÔN HỌC TRỰC TUYẾN

1. Truy cập hệ thống đăng ký môn học trong thời gian quy định.
2. Chọn môn học theo khung chương trình đào tạo.
3. Xác nhận đăng ký và thanh toán học phí (nếu có).
4. Sau khi đăng ký thành công, hệ thống sẽ gửi xác nhận qua email.""",
        "adversarial_title": "Hướng dẫn không cần đăng ký môn học",
        "adversarial_content": """HƯỚNG DẪN KHÔNG CẦN ĐĂNG KÝ MÔN HỌC

Học viện không yêu cầu đăng ký môn học.
Học viên tự do chọn môn học và tham gia lớp bất kỳ.
Không cần xác nhận hay thanh toán.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Phúc khảo điểm",
        "normal_title": "Quy trình phúc khảo điểm",
        "normal_content": """QUY TRÌNH PHÚC KHẢO ĐIỂM

1. Học viên nộp đơn xin phúc khảo trong vòng 7 ngày kể từ khi công bố điểm.
2. Đơn phúc khảo được gửi về Phòng Đào tạo.
3. Hội đồng phúc khảo sẽ xem xét và ra quyết định trong vòng 10 ngày làm việc.
4. Kết quả phúc khảo được thông báo bằng văn bản.""",
        "adversarial_title": "Quy trình không phúc khảo điểm",
        "adversarial_content": """QUY TRÌNH KHÔNG PHÚC KHẢO ĐIỂM

Điều 1: Học viên không được phép phúc khảo điểm.
Điều 2: Điểm số là quyết định cuối cùng của giảng viên.
Điều 3: Không có cơ chế khiếu nại về điểm số.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Thực tập tốt nghiệp",
        "normal_title": "Quy định thực tập tốt nghiệp",
        "normal_content": """QUY ĐỊNH THỰC TẬP TỐT NGHIỆP

Điều 1: Thực tập tốt nghiệp là bắt buộc đối với tất cả học viên.
Điều 2: Thời gian thực tập tối thiểu 2 tháng tại doanh nghiệp hoặc tổ chức.
Điều 3: Học viên phải nộp báo cáo thực tập và được đánh giá bởi giảng viên hướng dẫn.
Điều 4: Kết quả thực tập là điều kiện để xét tốt nghiệp.""",
        "adversarial_title": "Quy định miễn thực tập",
        "adversarial_content": """QUY ĐỊNH MIỄN THỰC TẬP

Điều 1: Học viên được miễn thực tập tốt nghiệp.
Điều 2: Không yêu cầu thời gian thực tập hay báo cáo.
Điều 3: Tốt nghiệp không phụ thuộc vào thực tập.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Luận văn tốt nghiệp",
        "normal_title": "Quy định luận văn tốt nghiệp",
        "normal_content": """QUY ĐỊNH LUẬN VĂN TỐT NGHIỆP

Điều 1: Học viên phải hoàn thành luận văn tốt nghiệp để được công nhận tốt nghiệp.
Điều 2: Luận văn phải được hướng dẫn bởi giảng viên có chuyên môn.
Điều 3: Thời gian thực hiện luận văn tối thiểu 3 tháng.
Điều 4: Luận văn được đánh giá bởi Hội đồng chấm luận văn.""",
        "adversarial_title": "Quy định miễn luận văn",
        "adversarial_content": """QUY ĐỊNH MIỄN LUẬN VĂN

Điều 1: Học viên được miễn làm luận văn tốt nghiệp.
Điều 2: Tốt nghiệp dựa trên điểm tích lũy và hoạt động khác.
Điều 3: Không yêu cầu bảo vệ luận văn.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Đánh giá giảng viên",
        "normal_title": "Quy định đánh giá giảng viên",
        "normal_content": """QUY ĐỊNH ĐÁNH GIÁ GIẢNG VIÊN

Điều 1: Học viên có trách nhiệm đánh giá giảng viên sau mỗi học kỳ.
Điều 2: Đánh giá được thực hiện qua hệ thống trực tuyến.
Điều 3: Kết quả đánh giá được sử dụng để cải thiện chất lượng giảng dạy.
Điều 4: Thông tin đánh giá được bảo mật.""",
        "adversarial_title": "Quy định không đánh giá giảng viên",
        "adversarial_content": """QUY ĐỊNH KHÔNG ĐÁNH GIÁ GIẢNG VIÊN

Điều 1: Học viên không cần đánh giá giảng viên.
Điều 2: Không có hệ thống đánh giá giảng dạy.
Điều 3: Chất lượng giảng dạy không được giám sát.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Sử dụng cơ sở vật chất",
        "normal_title": "Quy định sử dụng cơ sở vật chất",
        "normal_content": """QUY ĐỊNH SỬ DỤNG CƠ SỞ VẬT CHẤT

1. Học viên sử dụng thiết bị, phòng học đúng mục đích.
2. Không tự ý di chuyển hoặc tháo dỡ thiết bị.
3. Báo cáo ngay khi phát hiện hư hỏng.
4. Vi phạm sẽ bị xử lý theo quy định.""",
        "adversarial_title": "Quy định tự do sử dụng cơ sở vật chất",
        "adversarial_content": """QUY ĐỊNH TỰ DO SỬ DỤNG CƠ SỞ VẬT CHẤT

1. Học viên được tự do sử dụng mọi thiết bị, phòng học.
2. Không cần đăng ký hay xin phép.
3. Không chịu trách nhiệm về hư hỏng.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Thi lại",
        "normal_title": "Quy định thi lại",
        "normal_content": """QUY ĐỊNH THI LẠI

Điều 1: Học viên có thể đăng ký thi lại các môn chưa đạt.
Điều 2: Thời gian thi lại được tổ chức vào đầu mỗi học kỳ.
Điều 3: Điểm thi lại được tính vào điểm trung bình chung.
Điều 4: Học viên chỉ được đăng ký thi lại tối đa 2 lần cho một môn.""",
        "adversarial_title": "Quy định không thi lại",
        "adversarial_content": """QUY ĐỊNH KHÔNG THI LẠI

Điều 1: Học viên không được phép thi lại bất kỳ môn nào.
Điều 2: Điểm thi là kết quả cuối cùng.
Điều 3: Không có cơ hội cải thiện điểm.""",
        "adversarial_type": "conflicting"
    },
    # Textbook-based group
    {
        "topic": "Nội dung giáo trình",
        "normal_title": "Giáo trình Toán cao cấp - Chương 1: Giới hạn",
        "normal_content": """CHƯƠNG 1: GIỚI HẠN VÀ LIÊN TỤC

Giới hạn của hàm số là khái niệm cơ bản của giải tích.
Định nghĩa: Hàm số f(x) có giới hạn L khi x tiến tới a nếu với mọi ε > 0, tồn tại δ > 0 sao cho |f(x)-L| < ε khi 0 < |x-a| < δ.
Hàm số liên tục tại a nếu giới hạn của f(x) tại a bằng f(a).""",
        "adversarial_title": "Giáo trình Toán cao cấp - Chương 1: Giới hạn (phiên bản sai)",
        "adversarial_content": """CHƯƠNG 1: GIỚI HẠN (PHIÊN BẢN KHÔNG CHÍNH XÁC)

Giới hạn của hàm số là giá trị của hàm tại điểm đó.
Hàm số liên tục tại a nếu giới hạn trái và phải không bằng nhau.
Đây là định nghĩa chính thức được sử dụng trong giáo trình.""",
        "adversarial_type": "conflicting"
    },
    {
        "topic": "Bài giảng",
        "normal_title": "Bài giảng Hệ điều hành - Quản lý tiến trình",
        "normal_content": """QUẢN LÝ TIẾN TRÌNH

Tiến trình là chương trình đang thực thi.
PCB lưu thông tin trạng thái, thanh ghi, bộ đếm chương trình.
Các trạng thái: new, ready, running, blocked, terminated.
Điều phối CPU dùng các thuật toán FCFS, SJF, RR.""",
        "adversarial_title": "Bài giảng sai - Quản lý tiến trình",
        "adversarial_content": """QUẢN LÝ TIẾN TRÌNH (PHIÊN BẢN KHÔNG CHÍNH XÁC)

Tiến trình và luồng là một.
Không có trạng thái blocked.
Điều phối CPU dùng thuật toán ngẫu nhiên.""",
        "adversarial_type": "conflicting"
    },
]

# ============================================================
# CATEGORIES & TYPES
# ============================================================

CATEGORIES = [
    "Giáo trình",
    "Bài giảng",
    "Tài liệu tham khảo",
    "Quy chế",
    "Nghiên cứu",
    "Luận văn",
    "Đề cương",
    "Quy định",
    "Nội quy",
    "Thông báo",
    "Hướng dẫn",
    "Biên bản họp",
    "Đề thi",
    "Đáp án",
    "Ngân hàng câu hỏi",
    "Đề cương ôn tập",
]

SUBJECTS = list(SUBJECT_PARAGRAPHS.keys())
CHAPTER_TOPICS = []
for paras in SUBJECT_PARAGRAPHS.values():
    CHAPTER_TOPICS.extend(paras.keys())
CHAPTER_TOPICS = list(set(CHAPTER_TOPICS))

DEPARTMENTS = [
    {"code": "K_CNTT", "name": "Khoa Công nghệ thông tin"},
    {"code": "K_TOAN", "name": "Khoa Toán"},
    {"code": "K_LY", "name": "Khoa Vật lý"},
    {"code": "K_HOA", "name": "Khoa Hóa học"},
    {"code": "K_NN", "name": "Khoa Ngoại ngữ"},
    {"code": "K_QS", "name": "Khoa Quân sự"},
    {"code": "P_DAOTAO", "name": "Phòng Đào tạo"},
    {"code": "P_KT", "name": "Phòng Khảo thí"},
    {"code": "P_KHCN", "name": "Phòng Khoa học Công nghệ"},
    {"code": "P_CHINHTRI", "name": "Phòng Chính trị"},
]

SECURITY_LEVELS = ['public', 'internal', 'restricted', 'confidential']
SCOPE_TYPES = ['all', 'role', 'department', 'custom']

USER_IDS = ['user-001', 'user-002', 'user-003', 'user-004', 'user-005']
ROLE_CODES = ['ADMIN', 'BGD', 'P2', 'GV', 'HV', 'KT']

# ============================================================
# COHERENT CONTENT GENERATORS
# ============================================================

def get_paragraphs_for_subject_topic(subject, topic):
    """Return a list of paragraphs for a given subject and topic."""
    if subject in SUBJECT_PARAGRAPHS:
        if topic in SUBJECT_PARAGRAPHS[subject]:
            return SUBJECT_PARAGRAPHS[subject][topic]
    return GENERAL_PARAGRAPHS

def generate_coherent_chapter(subject, topic, num_paragraphs=None):
    """Generate a coherent chapter by selecting paragraphs from the database."""
    paras = get_paragraphs_for_subject_topic(subject, topic)
    if not paras:
        return "Nội dung không có sẵn."
    if num_paragraphs is None:
        num_paragraphs = random.randint(3, 6)
    selected = random.sample(paras, min(num_paragraphs, len(paras)))
    result = []
    for i, para in enumerate(selected):
        if i == 0:
            result.append(para)
        else:
            if random.random() < 0.3:
                transitions = ["Tiếp theo, ", "Mặt khác, ", "Đáng chú ý là, ", "Ngoài ra, "]
                result.append(random.choice(transitions) + para[0].lower() + para[1:])
            else:
                result.append(para)
    return "\n\n".join(result)

def generate_chapter_content(num_chapters=3):
    """Generate content for a textbook-like document with multiple chapters."""
    subject = random.choice(SUBJECTS)
    available_topics = list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())
    if not available_topics:
        available_topics = ["Tổng quan"]
    chosen_topics = random.sample(available_topics, min(num_chapters, len(available_topics)))
    content = []
    for i, topic in enumerate(chosen_topics, 1):
        content.append(f"Chương {i}: {topic}")
        content.append(generate_coherent_chapter(subject, topic, random.randint(3, 5)))
    return "\n\n".join(content)

def generate_quy_dinh_content(topics=None):
    """Generate a regulation document with coherent articles."""
    if topics is None:
        topics = [
            "Quy định về đào tạo",
            "Quy định về thi cử",
            "Quy định về tốt nghiệp",
            "Quy định về học bổng",
            "Quy định về kỷ luật học sinh",
            "Quy định về luận văn tốt nghiệp",
            "Quy định về thực tập",
            "Quy định về chuyển ngành",
            "Quy định về bảo lưu kết quả",
            "Quy định về văn bằng chứng chỉ",
        ]
    chosen = random.sample(topics, random.randint(3, 6))
    lines = ["QUY ĐỊNH CỦA HỌC VIỆN\n"]
    for i, topic in enumerate(chosen, 1):
        lines.append(f"\nĐiều {i}: {topic}")
        if topic.startswith("Quy định về"):
            keyword = topic.replace("Quy định về", "").strip()
            para = f"  Học viện ban hành quy định cụ thể về {keyword} nhằm đảm bảo tính thống nhất và công bằng cho tất cả học viên. Các nội dung chi tiết được trình bày trong các khoản dưới đây."
        else:
            para = f"  {generate_coherent_chapter('Giáo trình', 'Tổng quan', 1)}"
        lines.append(para)
    return "\n".join(lines)

def generate_research_content():
    """Generate a research paper with coherent sections."""
    subject = random.choice(SUBJECTS)
    title = f"Nghiên cứu về {subject} trong giáo dục"
    abstract = generate_coherent_chapter(subject, "Tổng quan", 2)
    intro = generate_coherent_chapter(subject, random.choice(list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())), 3)
    method = generate_coherent_chapter(subject, random.choice(list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())), 2)
    results = generate_coherent_chapter(subject, random.choice(list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())), 3)
    conclusion = generate_coherent_chapter(subject, "Tổng quan", 2)
    return f"""
    {title}
    
    TÓM TẮT
    {abstract}
    
    1. ĐẶT VẤN ĐỀ
    {intro}
    
    2. PHƯƠNG PHÁP NGHIÊN CỨU
    {method}
    
    3. KẾT QUẢ VÀ THẢO LUẬN
    {results}
    
    4. KẾT LUẬN
    {conclusion}
    
    TÀI LIỆU THAM KHẢO
    1. Tài liệu tham khảo 1
    2. Tài liệu tham khảo 2
    3. Tài liệu tham khảo 3
    """

def generate_luan_van_content():
    """Generate a thesis with coherent chapters."""
    subject = random.choice(SUBJECTS)
    title = f"Luận văn về {subject} trong đào tạo quân sự"
    intro = generate_coherent_chapter(subject, "Tổng quan", 3)
    theory = generate_coherent_chapter(subject, random.choice(list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())), 4)
    method = generate_coherent_chapter(subject, random.choice(list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())), 3)
    results = generate_coherent_chapter(subject, random.choice(list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())), 4)
    conclusion = generate_coherent_chapter(subject, "Tổng quan", 2)
    return f"""
    {title}
    
    TÓM TẮT
    {generate_coherent_chapter(subject, "Tổng quan", 2)}
    
    Chương 1: GIỚI THIỆU
    {intro}
    
    Chương 2: CƠ SỞ LÝ THUYẾT
    {theory}
    
    Chương 3: PHƯƠNG PHÁP NGHIÊN CỨU
    {method}
    
    Chương 4: KẾT QUẢ VÀ THẢO LUẬN
    {results}
    
    Chương 5: KẾT LUẬN VÀ KIẾN NGHỊ
    {conclusion}
    
    TÀI LIỆU THAM KHẢO
    1. Reference 1
    2. Reference 2
    3. Reference 3
    4. Reference 4
    """

# ----- EXAM CONTENT -----
def generate_exam_paper_content(subject=None, exam_type="midterm"):
    """Generate a coherent exam paper with questions related to the subject."""
    if subject is None:
        subject = random.choice(SUBJECTS)
    topics = list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())
    if not topics:
        topics = ["Tổng quan"]
    selected_topics = random.sample(topics, min(3, len(topics)))
    
    mc_questions = []
    for i in range(1, 11):
        topic = random.choice(selected_topics)
        stem = f"Câu {i}: Về {topic.lower()}, vấn đề nào sau đây là đúng?"
        options = [
            f"A. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}",
            f"B. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}",
            f"C. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}",
            f"D. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}"
        ]
        mc_questions.append(f"{stem}\n   " + "\n   ".join(options))
    
    tf_questions = []
    for i in range(1, 6):
        topic = random.choice(selected_topics)
        tf_questions.append(f"Câu {i}: {generate_coherent_chapter(subject, topic, 1).split('.')[0]}? (Đúng/Sai)")
    
    sa_questions = []
    for i in range(1, random.randint(2, 4)):
        topic = random.choice(selected_topics)
        sa_questions.append(f"Câu {i}: Trình bày ngắn gọn về {topic.lower()}.")
    
    content = f"""
ĐỀ THI MÔN: {subject}
{exam_type.upper()} - {random.choice(['Học kỳ 1', 'Học kỳ 2', 'Học kỳ hè'])} NĂM HỌC {random.randint(2024, 2026)}

Thời gian: 90 phút
Số trang: {random.randint(2, 4)}

--- HƯỚNG DẪN ---
- Thí sinh đọc kỹ đề trước khi làm.
- Không được sử dụng tài liệu.
- Nộp bài sau khi hết giờ.

---
PHẦN I: TRẮC NGHIỆM (4 điểm)
{chr(10).join(mc_questions)}

PHẦN II: ĐÚNG/SAI (2 điểm)
{chr(10).join(tf_questions)}

PHẦN III: TỰ LUẬN (4 điểm)
{chr(10).join(sa_questions)}
"""
    return content

def generate_answer_key_content(exam_subject=None, exam_id=None):
    """Generate an answer key with coherent explanations."""
    if exam_subject is None:
        exam_subject = random.choice(SUBJECTS)
    mc_answers = [random.choice(['A', 'B', 'C', 'D']) for _ in range(10)]
    tf_answers = [random.choice(['Đúng', 'Sai']) for _ in range(5)]
    sa_hints = [f"Gợi ý: {generate_coherent_chapter(exam_subject, random.choice(list(SUBJECT_PARAGRAPHS.get(exam_subject, {}).keys())), 1)}" for _ in range(random.randint(2, 3))]
    content = f"""
ĐÁP ÁN ĐỀ THI MÔN: {exam_subject}
{'(Đề số ' + str(exam_id) + ')' if exam_id else ''}

--- PHẦN I: TRẮC NGHIỆM ---
{chr(10).join([f"Câu {i+1}: {mc_answers[i]}" for i in range(10)])}

--- PHẦN II: ĐÚNG/SAI ---
{chr(10).join([f"Câu {i+1}: {tf_answers[i]}" for i in range(5)])}

--- PHẦN III: TỰ LUẬN (GỢI Ý) ---
{chr(10).join([f"Câu {i+1}: {sa_hints[i]}" for i in range(len(sa_hints))])}
"""
    return content

def generate_question_bank_content(subject=None, num_questions=20):
    """Generate a question bank with coherent questions."""
    if subject is None:
        subject = random.choice(SUBJECTS)
    topics = list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())
    if not topics:
        topics = ["Tổng quan"]
    questions = []
    for i in range(1, num_questions+1):
        difficulty = random.choice(["Dễ", "Trung bình", "Khó"])
        q_type = random.choice(["Trắc nghiệm", "Tự luận", "Đúng/Sai"])
        topic = random.choice(topics)
        if q_type == "Trắc nghiệm":
            stem = f"Về {topic.lower()}, phát biểu nào sau đây là chính xác?"
            options = [
                f"A. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}",
                f"B. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}",
                f"C. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}",
                f"D. {generate_coherent_chapter(subject, topic, 1).split('.')[0][:60]}"
            ]
            correct = random.choice(['A', 'B', 'C', 'D'])
            q_text = f"{stem}\n   " + "\n   ".join(options)
        elif q_type == "Tự luận":
            q_text = f"Phân tích ý nghĩa của {topic.lower()} trong {subject}."
            correct = "Đáp án mở"
        else:  # True/False
            q_text = f"{generate_coherent_chapter(subject, topic, 1).split('.')[0]}? (Đúng/Sai)"
            correct = random.choice(["Đúng", "Sai"])
        questions.append(f"""
[Q{i}] {q_type} - {difficulty}
Câu hỏi: {q_text}
Đáp án: {correct}
Điểm: {random.choice([0.5, 1.0, 1.5, 2.0])}
Chủ đề: {topic}
""")
    content = f"""
NGÂN HÀNG CÂU HỎI MÔN: {subject}
Tổng số câu: {num_questions}

""" + "\n".join(questions)
    return content

def generate_de_cuong_on_tap(subject=None):
    """Generate a revision outline with coherent content."""
    if subject is None:
        subject = random.choice(SUBJECTS)
    topics = list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())
    if not topics:
        topics = ["Tổng quan"]
    chosen_topics = random.sample(topics, min(4, len(topics)))
    content = f"""
ĐỀ CƯƠNG ÔN TẬP MÔN: {subject}
Học kỳ: {random.choice(['Học kỳ 1', 'Học kỳ 2'])}
Năm học: {random.randint(2024, 2026)}

I. CẤU TRÚC ĐỀ THI
- Hình thức: {random.choice(['Trắc nghiệm + Tự luận', 'Tự luận', 'Trắc nghiệm'])}
- Thời gian: {random.choice([60, 90, 120])} phút
- Tổng điểm: 10

II. NỘI DUNG ÔN TẬP
{chr(10).join([f"- {topic}" for topic in chosen_topics])}

III. DẠNG BÀI TẬP
{chr(10).join([f"- Dạng {i+1}: {generate_coherent_chapter(subject, topic, 1)}" for i, topic in enumerate(chosen_topics[:3])])}

IV. TÀI LIỆU THAM KHẢO
- Giáo trình chính
- Bài giảng trên lớp
- Đề thi các năm trước
"""
    return content

# ----- ADMINISTRATIVE / OPERATIONAL DOCUMENTS -----

def generate_thong_bao_content():
    """Generate a realistic announcement (Thông báo)."""
    topics = [
        "lịch thi giữa kỳ học kỳ 1 năm học 2025-2026",
        "thu học phí học kỳ 2 năm học 2024-2025",
        "tuyển sinh đào tạo từ xa năm 2026",
        "lịch nghỉ Tết Nguyên Đán 2026",
        "hội thảo khoa học 'AI trong Giáo dục'",
        "bảo vệ luận văn tốt nghiệp đợt 1/2026",
        "lễ khai giảng năm học 2026-2027",
        "tổ chức thi cuối kỳ trực tuyến",
        "chương trình học bổng trao đổi sinh viên",
        "nộp hồ sơ đăng ký thực tập doanh nghiệp"
    ]
    topic = random.choice(topics)
    so_hieu = random.randint(100, 999)
    ngay_ban_hanh = datetime.now().strftime("%d/%m/%Y")
    
    content = f"""THÔNG BÁO SỐ {so_hieu}/TB-HV

Kính gửi: Toàn thể cán bộ, giảng viên và học viên Học viện

Căn cứ kế hoạch đào tạo và hoạt động của Học viện năm học 2025-2026;

Học viện thông báo về việc {topic} như sau:

1. Nội dung:
{random.choice([
    f"Học viện tổ chức {topic} theo đúng quy định hiện hành.",
    f"Để đảm bảo tiến độ công việc, Học viện yêu cầu các đơn vị triển khai {topic} đúng thời gian quy định.",
    f"Nhằm nâng cao chất lượng đào tạo, Học viện triển khai {topic} với các nội dung chi tiết đính kèm."
])}

2. Thời gian thực hiện:
- Bắt đầu: {random.randint(1, 30)}/{random.randint(1, 12)}/{random.randint(2025, 2026)}
- Kết thúc: {random.randint(1, 28)}/{random.randint(1, 12)}/{random.randint(2025, 2026)}

3. Trách nhiệm:
Các đơn vị, phòng ban và cá nhân có liên quan chịu trách nhiệm tổ chức thực hiện nghiêm túc, báo cáo kết quả về Phòng Đào tạo trước ngày {random.randint(1, 28)}/{random.randint(1, 12)}/{random.randint(2025, 2026)}.

4. Địa điểm / Hình thức:
{random.choice([
    "Hội trường A - Học viện",
    "Phòng họp trực tuyến (Zoom / Microsoft Teams)",
    "Cổng thông tin điện tử của Học viện"
])}

Nơi nhận:
- Ban Giám đốc (để b/c);
- Các phòng, khoa (để t/h);
- Lưu VT.

TM. HỌC VIỆN
GIÁM ĐỐC
(Ký, đóng dấu)
"""
    return content

def generate_huong_dan_content():
    """Generate realistic guidelines (Hướng dẫn)."""
    topics = [
        "đăng ký môn học trực tuyến",
        "thực tập tốt nghiệp",
        "làm luận văn tốt nghiệp",
        "sử dụng hệ thống quản lý đào tạo (LMS)",
        "nộp hồ sơ xét tốt nghiệp",
        "đăng ký thi lại / học lại",
        "chuyển đổi ngành học",
        "xin cấp bảng điểm / văn bằng"
    ]
    topic = random.choice(topics)
    so_hieu = random.randint(200, 999)
    ngay_ban_hanh = datetime.now().strftime("%d/%m/%Y")
    
    steps = [
        f"Bước 1: Chuẩn bị hồ sơ theo danh mục yêu cầu. Danh mục hồ sơ bao gồm: đơn xin {topic}, bản sao các giấy tờ liên quan, và các minh chứng kèm theo.",
        f"Bước 2: Nộp hồ sơ tại Phòng {random.choice(['Đào tạo', 'Khảo thí', 'Chính trị'])} hoặc qua cổng dịch vụ công trực tuyến của Học viện.",
        f"Bước 3: Nhận giấy hẹn hoặc thông báo kết quả qua email / tin nhắn. Thời gian xử lý tối đa là {random.randint(3, 7)} ngày làm việc.",
        f"Bước 4: Kiểm tra kết quả trên hệ thống và liên hệ bộ phận hỗ trợ nếu có vướng mắc."
    ]
    random.shuffle(steps)
    
    content = f"""HƯỚNG DẪN SỐ {so_hieu}/HD-HV

Hướng dẫn về việc {topic}

Thực hiện theo yêu cầu của Ban Giám đốc, Học viện hướng dẫn quy trình {topic} như sau:

I. ĐỐI TƯỢNG ÁP DỤNG
Tất cả học viên đang theo học tại Học viện có nhu cầu {topic}.

II. QUY TRÌNH THỰC HIỆN
{chr(10).join(steps)}

III. LƯU Ý
- Hồ sơ phải được nộp đúng mẫu và đầy đủ chữ ký, đóng dấu (nếu có).
- Hạn nộp hồ sơ: trước ngày {random.randint(1, 28)}/{random.randint(1, 12)}/{random.randint(2025, 2026)}.
- Mọi thắc mắc xin liên hệ Phòng {random.choice(['Đào tạo', 'Khảo thí'])} (số điện thoại: 024.3.XXX.XXX).

Nơi nhận:
- Như trên;
- Lưu VT.

TM. HỌC VIỆN
TRƯỞNG PHÒNG ĐÀO TẠO
(Ký, đóng dấu)
"""
    return content

def generate_bien_ban_hop_content():
    """Generate realistic meeting minutes (Biên bản họp)."""
    chủ_trì = random.choice([
        "TS. Nguyễn Văn A - Giám đốc Học viện",
        "TS. Nguyễn Văn B - Phó Giám đốc phụ trách Đào tạo",
        "TS. Trần Văn C - Trưởng phòng Đào tạo"
    ])
    tham_gia = random.sample([
        "PGS.TS. Lê Thị D - Trưởng khoa CNTT",
        "TS. Phạm Văn E - Trưởng khoa Toán",
        "ThS. Nguyễn Thị F - Phó phòng Khảo thí",
        "ThS. Hoàng Văn G - Trưởng bộ môn Hệ điều hành",
        "CN. Lê Văn H - Đại diện sinh viên",
        "TS. Vũ Thị I - Đại diện Phòng Khoa học Công nghệ"
    ], random.randint(3, 5))
    
    noi_dung = f"""
{random.choice([
    "1. Đánh giá kết quả đào tạo học kỳ 1 năm học 2025-2026.",
    "1. Triển khai kế hoạch khảo thí học kỳ 2.",
    "1. Rà soát và cập nhật chương trình đào tạo ngành Công nghệ thông tin.",
    "1. Bàn bạc và thống nhất kế hoạch tuyển sinh năm học 2026-2027.",
    "1. Đánh giá tình hình thực hiện đề tài nghiên cứu khoa học cấp cơ sở."
])}
2. {generate_coherent_chapter(random.choice(SUBJECTS), random.choice(CHAPTER_TOPICS), 2)}
3. Công tác chuẩn bị cho đợt kiểm định chất lượng giáo dục.
4. Các vấn đề phát sinh và kiến nghị.
"""
    
    ket_luan = f"""
1. {random.choice([
    "Thông qua kế hoạch đào tạo và lịch thi mới.",
    "Nhất trí thông qua dự thảo quy chế đào tạo.",
    "Thống nhất phương án tuyển sinh và chỉ tiêu.",
    "Đề nghị các đơn vị báo cáo tiến độ trước ngày 30/06/2026."
])}
2. {random.choice([
    "Hoàn thành báo cáo tổng kết trước ngày 15/07/2026.",
    "Tổ chức tập huấn cho giảng viên về phương pháp giảng dạy mới.",
    "Phân công ThS. Nguyễn Văn X phụ trách soạn thảo văn bản hướng dẫn."
])}
"""
    
    content = f"""BIÊN BẢN HỌP
Số {random.randint(1, 99)}/BB-HV
Ngày {datetime.now().strftime('%d/%m/%Y')}

I. THÀNH PHẦN
- Chủ trì: {chủ_trì}
- Thư ký: ThS. {random.choice(['Nguyễn Thị', 'Lê Văn', 'Trần Văn'])} {random.choice(['A', 'B', 'C'])}
- Tham dự: {', '.join(tham_gia)}

II. NỘI DUNG HỌP
{noi_dung}

III. KẾT LUẬN VÀ KIẾN NGHỊ
{ket_luan}

IV. Ý KIẾN KHÁC (nếu có)
{random.choice(['Không có ý kiến khác.', 'Có ý kiến về việc điều chỉnh thời gian thi môn Cấu trúc dữ liệu.'])}
"""
    return content

def generate_de_cuong_content():
    """Generate a concise course outline (Đề cương môn học)."""
    subject = random.choice(SUBJECTS)
    topics = list(SUBJECT_PARAGRAPHS.get(subject, {}).keys())
    if not topics:
        topics = ["Tổng quan"]
    selected_topics = random.sample(topics, min(4, len(topics)))
    content = f"""ĐỀ CƯƠNG MÔN HỌC {subject}

I. THÔNG TIN CHUNG
- Số tín chỉ: {random.randint(2, 4)}
- Số tiết: {random.randint(30, 60)}
- Điều kiện tiên quyết: {random.choice(['Không', 'Toán cao cấp', 'Lập trình cơ bản', 'Không yêu cầu'])}
- Đánh giá: {random.choice(['Thi cuối kỳ (60%) + Bài tập (40%)', 'Thi giữa kỳ (30%) + Cuối kỳ (70%)', 'Đồ án (100%)'])}

II. MỤC TIÊU
- Trang bị cho học viên kiến thức nền tảng về {subject}.
- Phát triển kỹ năng phân tích và giải quyết vấn đề trong lĩnh vực {subject}.

III. NỘI DUNG CHI TIẾT
{chr(10).join([f"- Chương {i+1}: {topic}" for i, topic in enumerate(selected_topics)])}

IV. TÀI LIỆU THAM KHẢO
1. {subject} - Giáo trình chính, NXB Giáo dục.
2. Bài giảng trên hệ thống LMS.
3. Tài liệu tham khảo bổ sung theo hướng dẫn của giảng viên.
"""
    return content

# ============================================================
# GROUPED DOCUMENT GENERATION (2 normal + 1 adversarial)
# ============================================================

def generate_group_docs(group: dict, base_doc_id: int, dept: dict) -> list[dict]:
    """
    Generate 2 normal + 1 adversarial docs for a topic group.
    Returns a list of 3 doc_info dicts.
    """
    # Choose file types (distribute)
    file_types = random.choices(FILE_TYPES, weights=FILE_TYPE_PROBS, k=3)
    # Ensure we don't have duplicate file types for variety, but we can allow.
    
    docs = []
    
    # Normal doc 1
    normal1 = {
        "doc_id": f"DOC-{base_doc_id:04d}",
        "title": group["normal_title"],
        "category": "Quy định",  # or appropriate
        "department_code": dept["code"],
        "department_name": dept["name"],
        "file_type": file_types[0],
        "subject": "",
        "content": group["normal_content"],
        "is_adversarial": False,
        "adversarial_type": "none",
        "conflict_group_id": group["topic"],
        "security_level": "internal",  # authoritative
        "source_system": "official",
    }
    docs.append(normal1)
    
    # Normal doc 2 (slightly different title but same content, or duplicate? We can use same content or slight variant)
    normal2_title = group["normal_title"] + " (bản chi tiết)"
    normal2_content = group["normal_content"] + "\n\n(Đính kèm phụ lục hướng dẫn thực hiện.)"
    normal2 = {
        "doc_id": f"DOC-{base_doc_id+1:04d}",
        "title": normal2_title,
        "category": "Quy định",
        "department_code": dept["code"],
        "department_name": dept["name"],
        "file_type": file_types[1],
        "subject": "",
        "content": normal2_content,
        "is_adversarial": False,
        "adversarial_type": "none",
        "conflict_group_id": group["topic"],
        "security_level": "confidential",
        "source_system": "official",
    }
    docs.append(normal2)
    
    # Adversarial doc
    adv = {
        "doc_id": f"DOC-{base_doc_id+2:04d}",
        "title": group["adversarial_title"],
        "category": "Quy định",  # or appropriate
        "department_code": dept["code"],
        "department_name": dept["name"],
        "file_type": file_types[2],
        "subject": "",
        "content": group["adversarial_content"],
        "is_adversarial": True,
        "adversarial_type": group.get("adversarial_type", "conflicting"),
        "conflict_group_id": group["topic"],
        "security_level": "public",  # low trust
        "source_system": "manual_upload",
    }
    docs.append(adv)
    
    return docs

# ============================================================
# DOCUMENT GENERATORS (unchanged)
# ============================================================

def generate_pdf_text(doc_info, output_dir):
    """Generate normal PDF using reportlab"""
    try:
        pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
        font_name = 'Arial'
    except:
        font_name = 'Helvetica'

    subdir = 'normal'
    if doc_info.get('is_adversarial', False):
        subdir = 'adversarial'
    
    pdf_filename = f"{doc_info['doc_id']}.pdf"
    pdf_path = os.path.join(output_dir, subdir, pdf_filename)
    
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        alignment=TA_CENTER,
        fontSize=16,
        spaceAfter=12,
        fontName=font_name
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        fontName=font_name
    )
    
    story.append(Paragraph(escape(doc_info['title']), title_style))
    story.append(Spacer(1, 12))
    
    content = doc_info.get('content', generate_chapter_content(random.randint(3, 5)))
    
    for line in content.split('\n'):
        if line.strip():
            escaped_line = escape(line)
            story.append(Paragraph(escaped_line, body_style))
        else:
            story.append(Spacer(1, 6))
    
    doc.build(story)
    return pdf_path, pdf_filename


def generate_scanned_pdf(doc_info, output_dir):
    """
    Generate a scanned PDF (image-based) by rendering text onto a PIL image,
    splitting into pages at line boundaries (no mid-line cuts).
    """
    if not PIL_AVAILABLE:
        print(f"  Warning: PIL not available, falling back to normal PDF for {doc_info['doc_id']}")
        return generate_pdf_text(doc_info, output_dir)

    subdir = 'normal'
    if doc_info.get('is_adversarial', False):
        subdir = 'adversarial'
    
    pdf_filename = f"{doc_info['doc_id']}.pdf"
    pdf_path = os.path.join(output_dir, subdir, pdf_filename)
    
    # Prepare content
    content = doc_info.get('content', generate_chapter_content(random.randint(3, 5)))
    full_text = f"{doc_info['title']}\n\n{content}"
    
    # Parameters
    FONT_SIZE = 14
    IMAGE_WIDTH_PX = 600
    MARGIN = 40
    LINE_SPACING = 6

    # Load font
    try:
        font_paths = [
            'C:/Windows/Fonts/arial.ttf',
            '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
            '/System/Library/Fonts/Helvetica.ttc',
        ]
        font_path = None
        for path in font_paths:
            if os.path.exists(path):
                font_path = path
                break
        if font_path:
            font = ImageFont.truetype(font_path, FONT_SIZE)
        else:
            font = ImageFont.load_default()
    except:
        font = ImageFont.load_default()

    # Helper: wrap text
    def wrap_text(line, max_width, draw, font):
        words = line.split()
        if not words:
            return ['']
        wrapped = []
        current_line = []
        for word in words:
            test_line = ' '.join(current_line + [word])
            bbox = draw.textbbox((0,0), test_line, font=font)
            if bbox[2] - bbox[0] <= max_width:
                current_line.append(word)
            else:
                if current_line:
                    wrapped.append(' '.join(current_line))
                current_line = [word]
        if current_line:
            wrapped.append(' '.join(current_line))
        return wrapped

    # Prepare draw for measuring
    temp_img = PILImage.new('RGB', (1, 1), color='white')
    draw = ImageDraw.Draw(temp_img)
    max_text_width = IMAGE_WIDTH_PX - 2 * MARGIN

    # Process all lines
    raw_lines = full_text.split('\n')
    final_lines = []
    for raw_line in raw_lines:
        if raw_line.strip() == '':
            final_lines.append('')
        else:
            wrapped = wrap_text(raw_line, max_text_width, draw, font)
            final_lines.extend(wrapped)

    # Measure line heights and compute cumulative positions (including margins)
    line_heights = []
    cum_heights = [0]  # cum_heights[i] = y position of start of line i (0-based) relative to image top (including margin)
    y = MARGIN
    for line in final_lines:
        if line.strip() == '':
            h = FONT_SIZE + LINE_SPACING
        else:
            bbox = draw.textbbox((0,0), line, font=font)
            h = bbox[3] - bbox[1] or FONT_SIZE
        line_heights.append(h)
        cum_heights.append(y)  # start of this line
        y += h + LINE_SPACING
    total_height = y + MARGIN  # bottom margin

    # Render full image
    img = PILImage.new('RGB', (IMAGE_WIDTH_PX, total_height), color='white')
    draw = ImageDraw.Draw(img)
    y_pos = MARGIN
    for line, h in zip(final_lines, line_heights):
        if line.strip() != '':
            draw.text((MARGIN, y_pos), line, fill='black', font=font)
        y_pos += h + LINE_SPACING

    # Convert to PDF with pagination at line boundaries
    page_width_pt, page_height_pt = letter
    margin_pt = 0.75 * inch
    usable_width_pt = page_width_pt - 2 * margin_pt
    usable_height_pt = page_height_pt - 2 * margin_pt

    scale = usable_width_pt / IMAGE_WIDTH_PX
    page_pixel_height = usable_height_pt / scale

    c = canvas.Canvas(pdf_path, pagesize=letter)

    # Determine page breaks
    start_idx = 0  # index into final_lines (cum_heights index)
    while start_idx < len(final_lines):
        end_idx = start_idx
        while end_idx < len(final_lines):
            # check if line end fits
            if cum_heights[end_idx] + line_heights[end_idx] - cum_heights[start_idx] <= page_pixel_height:
                end_idx += 1
            else:
                break
        # If no line fits (shouldn't happen), take at least one line
        if end_idx == start_idx:
            end_idx = start_idx + 1

        y_start = cum_heights[start_idx]
        y_end = cum_heights[end_idx]  # start of next line, which is end of last line
        # Ensure we don't go beyond total_height
        if y_end > total_height:
            y_end = total_height
        # Crop
        slice_img = img.crop((0, y_start, IMAGE_WIDTH_PX, y_end))
        img_reader = ImageReader(slice_img)
        slice_height_pt = (y_end - y_start) * scale
        # Draw with top at margin_pt from top
        y_draw = page_height_pt - margin_pt - slice_height_pt
        c.drawImage(img_reader, margin_pt, y_draw,
                    width=usable_width_pt, height=slice_height_pt,
                    preserveAspectRatio=False)
        c.showPage()
        start_idx = end_idx

    c.save()
    return pdf_path, pdf_filename


def generate_docx(doc_info, output_dir):
    """Generate a DOCX file with proper heading styles"""
    doc = Document()
    
    title = doc.add_heading(doc_info['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Phân loại: {doc_info['category']}")
    doc.add_paragraph(f"Nguồn: {doc_info['department_name']}")
    doc.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("")
    
    content = doc_info.get('content', generate_chapter_content(random.randint(2, 4)))
    
    lines = content.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        
        if stripped.startswith(('Chương', 'Chuong', 'Phần', 'Phan', 'Part', 'Chapter')):
            if not re.search(r'(Chương|Chuong|Phần|Phan|Part|Chapter)\s+\d+', stripped, re.IGNORECASE):
                doc.add_paragraph(stripped)
            else:
                doc.add_heading(stripped, level=1)
        elif stripped.startswith(('Điều', 'Dieu', 'Section', 'Mục', 'Muc')):
            doc.add_heading(stripped, level=2)
        elif re.match(r'^\d+\.\d+\.?\s+', stripped):
            doc.add_heading(stripped, level=3)
        else:
            doc.add_paragraph(stripped)

    subdir = 'normal'
    if doc_info.get('is_adversarial', False):
        subdir = 'adversarial'
    
    filename = f"{doc_info['doc_id']}.docx"
    filepath = os.path.join(output_dir, subdir, filename)
    doc.save(filepath)
    return filepath, filename


def generate_txt(doc_info, output_dir):
    """Generate a TXT file"""
    lines = []
    lines.append("=" * 60)
    lines.append(doc_info["title"].center(60))
    lines.append("=" * 60)
    lines.append(f"Phân loại: {doc_info['category']}")
    lines.append(f"Nguồn: {doc_info['department_name']}")
    lines.append(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y')}")
    lines.append("=" * 60)
    lines.append("")
    
    content = doc_info.get('content', generate_chapter_content(random.randint(2, 4)))
    lines.append(content)

    subdir = 'normal'
    if doc_info.get('is_adversarial', False):
        subdir = 'adversarial'
    
    filename = f"{doc_info['doc_id']}.txt"
    filepath = os.path.join(output_dir, subdir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return filepath, filename


def generate_md(doc_info, output_dir):
    """Generate a Markdown (.md) file with proper Markdown formatting"""
    lines = []
    lines.append(f"# {doc_info['title']}\n")
    lines.append(f"**Phân loại:** {doc_info['category']}  ")
    lines.append(f"**Nguồn:** {doc_info['department_name']}  ")
    lines.append(f"**Ngày tạo:** {datetime.now().strftime('%d/%m/%Y')}\n")
    
    content = doc_info.get('content', generate_chapter_content(random.randint(2, 4)))
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            lines.append('')
            continue
        if stripped.startswith(('Chương', 'Chuong', 'Phần', 'Phan', 'Part', 'Chapter')):
            if re.search(r'(Chương|Chuong|Phần|Phan|Part|Chapter)\s+\d+', stripped, re.IGNORECASE):
                lines.append(f'## {stripped}')
            else:
                lines.append(stripped)
        elif stripped.startswith(('Điều', 'Dieu', 'Section', 'Mục', 'Muc')):
            lines.append(f'### {stripped}')
        elif re.match(r'^\d+\.\d+\.?\s+', stripped):
            lines.append(f'#### {stripped}')
        else:
            lines.append(f'  {stripped}')
    
    subdir = 'normal'
    if doc_info.get('is_adversarial', False):
        subdir = 'adversarial'
    
    filename = f"{doc_info['doc_id']}.md"
    filepath = os.path.join(output_dir, subdir, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return filepath, filename


# ============================================================
# GENERATE DOCUMENT INFO (for extra normal docs)
# ============================================================

def generate_extra_normal_doc_info(doc_id, dept=None):
    """Generate a normal document not part of conflict groups."""
    file_type = random.choices(FILE_TYPES, weights=FILE_TYPE_PROBS, k=1)[0]
    category = random.choice(CATEGORIES)
    if dept is None:
        dept = random.choice(DEPARTMENTS)
    
    if category in ["Giáo trình", "Tài liệu tham khảo", "Bài giảng"]:
        subject = random.choice(SUBJECTS)
        title = f"{subject} - {category} {random.randint(1, 3)}"
    elif category == "Đề cương":
        subject = random.choice(SUBJECTS)
        title = f"Đề cương môn học {subject}"
    elif category == "Quy chế":
        title = f"Quy chế đào tạo {random.randint(2024, 2026)}"
    elif category == "Nội quy":
        title = f"Nội quy học viên (Phiên bản {random.randint(1, 3)})"
    elif category == "Nghiên cứu":
        subject = random.choice(SUBJECTS)
        title = f"Nghiên cứu ứng dụng {subject} trong giáo dục"
    elif category == "Luận văn":
        subject = random.choice(SUBJECTS)
        title = f"Luận văn về {subject} trong đào tạo"
    elif category == "Thông báo":
        topics = [
            "lịch thi giữa kỳ", "thu học phí", "tuyển sinh", "lịch nghỉ",
            "hội thảo khoa học", "bảo vệ luận văn", "khai giảng", "thi cuối kỳ"
        ]
        title = f"Thông báo về {random.choice(topics)}"
    elif category == "Hướng dẫn":
        topics = ["đăng ký môn học", "thực tập", "tốt nghiệp", "sử dụng LMS", "thi lại"]
        title = f"Hướng dẫn {random.choice(topics)}"
    elif category == "Biên bản họp":
        title = f"Biên bản họp {random.choice(['khoa', 'hội đồng', 'đào tạo', 'khảo thí'])}"
    elif category == "Đề thi":
        subject = random.choice(SUBJECTS)
        title = f"Đề thi môn {subject} - {random.choice(['Giữa kỳ', 'Cuối kỳ', 'Thi thử'])}"
    elif category == "Đáp án":
        subject = random.choice(SUBJECTS)
        title = f"Đáp án đề thi môn {subject}"
    elif category == "Ngân hàng câu hỏi":
        subject = random.choice(SUBJECTS)
        title = f"Ngân hàng câu hỏi môn {subject}"
    elif category == "Đề cương ôn tập":
        subject = random.choice(SUBJECTS)
        title = f"Đề cương ôn tập môn {subject}"
    else:
        title = f"{category} {random.randint(1, 100)}"
    
    doc_info = {
        "doc_id": f"DOC-{doc_id:04d}",
        "title": title,
        "category": category,
        "department_code": dept["code"],
        "department_name": dept["name"],
        "file_type": file_type,
        "subject": subject if 'subject' in locals() else random.choice(SUBJECTS),
        "is_adversarial": False,
        "adversarial_type": "none",
        "conflict_group_id": None,
        "security_level": random.choice(['public', 'internal']),  # normal trust
        "source_system": "official",
    }
    
    # Generate content based on category
    subject = doc_info.get("subject", random.choice(SUBJECTS))
    if category in ["Quy chế", "Quy định", "Nội quy"]:
        content = generate_quy_dinh_content()
    elif category == "Nghiên cứu":
        content = generate_research_content()
    elif category == "Luận văn":
        content = generate_luan_van_content()
    elif category == "Đề thi":
        content = generate_exam_paper_content(subject, random.choice(["Giữa kỳ", "Cuối kỳ", "Thi thử"]))
    elif category == "Đáp án":
        content = generate_answer_key_content(subject, random.randint(100, 999))
    elif category == "Ngân hàng câu hỏi":
        content = generate_question_bank_content(subject, random.randint(15, 25))
    elif category == "Đề cương ôn tập":
        content = generate_de_cuong_on_tap(subject)
    elif category == "Đề cương":
        content = generate_de_cuong_content()
    elif category == "Thông báo":
        content = generate_thong_bao_content()
    elif category == "Hướng dẫn":
        content = generate_huong_dan_content()
    elif category == "Biên bản họp":
        content = generate_bien_ban_hop_content()
    else:
        content = generate_chapter_content(random.randint(2, 3))
    doc_info["content"] = content
    return doc_info

# ============================================================
# MONGODB SEED GENERATOR
# ============================================================

def generate_mongodb_seed(documents):
    """Generate MongoDB seed file (tai_lieu collection)"""
    now = datetime.now()
    now_str = now.isoformat()
    
    js_content = """// =====================================================
// MongoDB Seed Data - tai_lieu collection
// Auto-generated by scripts/generate_sample_docs.py
// =====================================================

if (db.documents.countDocuments() === 0) {
    db.documents.insertMany([
"""
    
    for i, doc_info in enumerate(documents):
        scope_type = random.choice(SCOPE_TYPES)
        
        access_role_codes = []
        access_department_codes = []
        access_user_ids = []
        
        if scope_type == "role":
            access_role_codes = random.sample(ROLE_CODES, random.randint(1, 3))
        elif scope_type == "department":
            access_department_codes = [doc_info['department_code']]
        elif scope_type == "custom":
            access_user_ids = random.sample(USER_IDS, random.randint(1, 3))
        
        subdir = 'normal'
        if doc_info.get('is_adversarial', False):
            subdir = 'adversarial'
        
        security_level = doc_info.get('security_level', random.choice(['public', 'internal']))
        
        file_type = doc_info['file_type']
        if file_type in ['pdf', 'scanned_pdf']:
            ext = 'pdf'
            mime_type = 'application/pdf'
        elif file_type == 'docx':
            ext = 'docx'
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif file_type == 'txt':
            ext = 'txt'
            mime_type = 'text/plain'
        elif file_type == 'md':
            ext = 'md'
            mime_type = 'text/markdown'
        else:
            ext = 'pdf'
            mime_type = 'application/pdf'
        
        # Build doc string
        doc_str = f"""      {{
        "docId": "{doc_info['doc_id']}",
        "title": "{doc_info['title']}",
        "category": "{doc_info['category']}",
        "originalName": "{doc_info['title'].replace(' ', '_')}.{ext}",
        "storedName": "{doc_info['doc_id']}.{ext}",
        "storagePath": "../../data/sample-docs/{subdir}/{doc_info['doc_id']}.{ext}",
        "mimeType": "{mime_type}",
        "size": {random.randint(100000, 5000000)},
        "securityLevel": "{security_level}",
        "scopeType": "{scope_type}",
        "accessRoleCodes": {json.dumps(access_role_codes)},
        "accessDepartmentCodes": {json.dumps(access_department_codes)},
        "accessUserIds": {json.dumps(access_user_ids)},
        "uploadedById": "{random.choice(USER_IDS)}",
        "uploadedByName": "Hệ thống (sample)",
        "createdAt": new Date("{now_str}"),
        "ingestStatus": "pending",
        "ingestStage": "queued",
        "ingestUpdatedAt": new Date("{now_str}"),
        "isSample": true,
        "isAdversarial": {str(doc_info.get('is_adversarial', False)).lower()},
        "adversarialType": "{doc_info.get('adversarial_type', 'none')}",
        "sourceSystem": "{doc_info.get('source_system', 'sample_offline')}",
        "isScanned": {str(file_type == 'scanned_pdf').lower()}
      }}"""
        
        if i < len(documents) - 1:
            js_content += doc_str + ",\n"
        else:
            js_content += doc_str + "\n"
    
    js_content += """   ]);
}

db.documents.createIndex({ createdAt: -1 });
db.documents.createIndex({ docId: 1 });
db.documents.createIndex({ isSample: 1 });
db.documents.createIndex({ ingestStatus: 1 });
db.documents.createIndex({ isAdversarial: 1 });
db.documents.createIndex({ adversarialType: 1 });
db.documents.createIndex({ isScanned: 1 });
"""
    
    seed_path = os.path.join(MONGODB_SEED_DIR, "03-seed-tai-lieu.js")
    with open(seed_path, "w", encoding="utf-8") as f:
        f.write(js_content)
    
    return seed_path

# ============================================================
# MAIN EXECUTION
# ============================================================

def main():
    print("=" * 60)
    print("Generating sample documents and MongoDB seed data...")
    print("=" * 60)
    
    num_adversarial = int(NUM_DOCUMENTS * ADVERSARIAL_RATIO)
    num_normal = NUM_DOCUMENTS - num_adversarial
    num_groups = num_adversarial  # each group has 1 adversarial
    normal_from_groups = num_groups * 2
    extra_normal = num_normal - normal_from_groups
    
    if extra_normal < 0:
        print(f"Warning: Not enough normal documents to form groups. Adjusting.")
        # Reduce groups to fit
        num_groups = min(num_groups, num_normal // 2)
        extra_normal = num_normal - num_groups * 2
        num_adversarial = num_groups  # adjust
    
    print(f"Generating {NUM_DOCUMENTS} documents:")
    print(f"   - Normal: {num_normal} (of which {normal_from_groups} from groups)")
    print(f"   - Adversarial: {num_adversarial} (in {num_groups} groups)")
    print(f"   - Extra normal docs: {extra_normal}")
    print("=" * 60)
    
    documents_meta = []
    doc_id = 1
    used_depts = set()
    
    # 1. Generate grouped triples
    for idx, group in enumerate(TOPIC_GROUPS[:num_groups]):
        # Pick a department (avoid repeating too much)
        dept = random.choice(DEPARTMENTS)
        # Generate 3 docs for this group
        group_docs = generate_group_docs(group, doc_id, dept)
        doc_id += 3
        documents_meta.extend(group_docs)
    
    # 2. Generate extra normal docs (non-conflict)
    for i in range(extra_normal):
        dept = random.choice(DEPARTMENTS)
        doc_info = generate_extra_normal_doc_info(doc_id, dept)
        doc_id += 1
        documents_meta.append(doc_info)
    
    # Write files
    for doc_info in documents_meta:
        if doc_info['file_type'] == 'scanned_pdf':
            filepath, filename = generate_scanned_pdf(doc_info, OUTPUT_DIR)
        elif doc_info['file_type'] == 'pdf':
            filepath, filename = generate_pdf_text(doc_info, OUTPUT_DIR)
        elif doc_info['file_type'] == 'docx':
            filepath, filename = generate_docx(doc_info, OUTPUT_DIR)
        elif doc_info['file_type'] == 'txt':
            filepath, filename = generate_txt(doc_info, OUTPUT_DIR)
        elif doc_info['file_type'] == 'md':
            filepath, filename = generate_md(doc_info, OUTPUT_DIR)
        else:
            continue
        
        status = "ADVERSARIAL" if doc_info.get('is_adversarial') else "Normal"
        group_info = f" [group: {doc_info.get('conflict_group_id', 'none')}]" if doc_info.get('conflict_group_id') else ""
        print(f"  {doc_info['doc_id']}: {filename} ({status}, {doc_info['file_type']}){group_info}")
    
    # Generate MongoDB seed
    seed_path = generate_mongodb_seed(documents_meta)
    print(f"\nMongoDB seed generated: {seed_path}")
    
    print("\n" + "=" * 60)
    print(f"Generated {len(documents_meta)} documents in '{OUTPUT_DIR}/'")
    print(f"Files location: {os.path.abspath(OUTPUT_DIR)}")
    print(f"MongoDB seed: {os.path.abspath(seed_path)}")
    print("=" * 60)
    print("\n Note: Adversarial documents are now grouped by topic (2 normal + 1 adversarial).")
    print("   Each group has `conflictGroupId` for tracing.")
    print("   Normal docs have `securityLevel: internal`, adversarial have `public`.")
    print("   This enables the demo to show authority selection.")
    print("\n File types distribution: scanned_pdf, docx, pdf, txt, md")
    print("\n Document categories now include:")
    print("   - Giáo trình, Bài giảng, Tài liệu tham khảo")
    print("   - Quy chế, Quy định, Nội quy")
    print("   - Nghiên cứu, Luận văn, Đề cương, Đề cương ôn tập")
    print("   - Đề thi, Đáp án, Ngân hàng câu hỏi")
    print("   - Thông báo, Hướng dẫn, Biên bản họp")

if __name__ == "__main__":
    main()

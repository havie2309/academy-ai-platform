from pathlib import Path


def extract_text(storage_path: str, mime_type: str = "") -> str:
    path = Path(storage_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {storage_path}")

    ext = path.suffix.lower()
    # File text thuần: đọc trực tiếp
    if ext in (".txt", ".md", ".csv"):
        return path.read_text(encoding="utf-8", errors="replace")

    # PDF: Dùng PyMuPDF 
    if ext == ".pdf" or "pdf" in mime_type:
        import fitz

        doc = fitz.open(str(path))
        try:
            return "\n".join(page.get_text() for page in doc).strip()
        finally:
            doc.close()

    # DOCX: dùng python-docx, lấy cả đoạn văn lẫn nội dung trong bảng
    if ext == ".docx" or "wordprocessingml" in mime_type:
        import docx

        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    
    raise ValueError(f"Định dạng chưa hỗ trợ ingest: {ext or mime_type}")
